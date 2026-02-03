import os
import sys
import pandas as pd
from docx import Document
import subprocess
import re

def convert_doc_to_docx(file_path):
    """使用 textutil 将 .doc 转换为 .docx"""
    temp_docx = file_path + "x"
    try:
        subprocess.run(["textutil", "-convert", "docx", file_path], check=True)
        return temp_docx
    except Exception as e:
        print(f"Error converting {file_path}: {e}")
        return None

def normalize(text):
    return "".join(text.split())

def extract_from_docx(file_path, station_name, file_name):
    """从 .docx 文件中提取表格数据"""
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return []

    data = []
    task_name = "未知任务"

    # 先寻找操作任务
    for table in doc.tables:
        for row in table.rows:
            text = " ".join(cell.text for cell in row.cells)
            text_norm = normalize(text)
            if "操作任务" in text_norm:
                # 尝试更宽泛的匹配
                match = re.search(r"操作任务[：:]\s*(.*)", text)
                if match:
                    task_name = match.group(1).split('\n')[0].strip()
                    if not task_name: # 如果冒号后面没内容，看后面一个单元格
                        for cell in row.cells:
                            c_text = cell.text.strip()
                            if c_text and "操作任务" not in c_text:
                                task_name = c_text
                                break
                    break
        if task_name != "未知任务":
            break
    
    if task_name == "未知任务":
        for para in doc.paragraphs:
            if "操作任务" in normalize(para.text):
                match = re.search(r"操作任务[：:]\s*(.*)", para.text)
                if match:
                    task_name = match.group(1).strip()
                    break

    # 提取步骤
    found_table = False
    for table in doc.tables:
        header_row_idx = -1
        order_idx = -1
        item_idx = -1
        time_idx = -1
        
        for i, row in enumerate(table.rows[:10]): # 假设表头在前10行
            row_texts = [normalize(c.text) for c in row.cells]
            if "序号" in row_texts or "顺序" in row_texts:
                header_row_idx = i
                for j, t in enumerate(row_texts):
                    if t in ["序号", "顺序"]:
                        order_idx = j
                    elif "操作项目" in t or "项目" in t:
                        if item_idx == -1: item_idx = j
                    elif "完成时间" in t or "时间" in t:
                        time_idx = j
                break
        
        if header_row_idx != -1 and order_idx != -1 and item_idx != -1:
            found_table = True
            for row in table.rows[header_row_idx + 1:]:
                cells = row.cells
                if len(cells) > max(order_idx, item_idx):
                    order = cells[order_idx].text.strip()
                    item = cells[item_idx].text.strip()
                    # 允许 1, 1.1, 1.1.1 等格式，或者只是数字
                    if order and (order[0].isdigit() or order == "序号"): 
                        if order == "序号": continue # 跳过可能的重复表头
                        finish_time = ""
                        if time_idx != -1 and time_idx < len(cells):
                            finish_time = cells[time_idx].text.strip()
                        
                        data.append({
                            "所属站场": station_name,
                            "文件名称": file_name,
                            "操作任务": task_name,
                            "顺序": order,
                            "操作项目": item,
                            "完成时间": finish_time
                        })

    if not found_table:
        # 段落解析增强型
        is_step_area = False
        last_order = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            
            norm_text = normalize(text)
            if "序号" in norm_text and ("操作项目" in norm_text or "项目" in norm_text):
                is_step_area = True
                continue
            
            if is_step_area:
                if any(x in text for x in ["备注：", "操作人：", "以下空白", "监护人："]):
                    is_step_area = False
                    continue
                
                # 模式1: "1 检查设备..."
                match = re.match(r"^(\d+(\.\d+)?)\s+(.*)", text)
                if match:
                    data.append({
                        "所属站场": station_name,
                        "文件名称": file_name,
                        "操作任务": task_name,
                        "顺序": match.group(1),
                        "操作项目": match.group(3),
                        "完成时间": ""
                    })
                    last_order = None
                # 模式2: 第一段是序号 "1", 第二段是内容 "检查设备..."
                elif text.isdigit() or re.match(r"^\d+(\.\d+)+$", text):
                    last_order = text
                elif last_order:
                    data.append({
                        "所属站场": station_name,
                        "文件名称": file_name,
                        "操作任务": task_name,
                        "顺序": last_order,
                        "操作项目": text,
                        "完成时间": ""
                    })
                    last_order = None

    return data

def main():
    root_dir = "/Users/yangjie/code/antigravity/关于开展典型操作票修编的通知"
    output_file = "/Users/yangjie/code/antigravity/操作票汇总表.xlsx"
    
    all_rows = []
    processed_files = 0
    
    for station in os.listdir(root_dir):
        station_path = os.path.join(root_dir, station)
        if not os.path.isdir(station_path):
            continue
        
        print(f"Processing station: {station}")
        
        for dirpath, dirnames, filenames in os.walk(station_path):
            for filename in filenames:
                if filename.startswith("~$") or filename == ".DS_Store": continue
                
                file_path = os.path.join(dirpath, filename)
                current_docx = None
                is_temp = False
                
                if filename.endswith(".docx"):
                    current_docx = file_path
                elif filename.endswith(".doc"):
                    current_docx = convert_doc_to_docx(file_path)
                    is_temp = True
                
                if current_docx:
                    processed_files += 1
                    rows = extract_from_docx(current_docx, station, filename)
                    all_rows.extend(rows)
                    
                    if is_temp and os.path.exists(current_docx):
                        os.remove(current_docx)

    if all_rows:
        df = pd.DataFrame(all_rows)
        # 清洗：移除没有操作项目的行
        df = df[df['操作项目'].str.strip() != ""]
        df.to_excel(output_file, index=False)
        print(f"Success! Consolidated {len(all_rows)} steps from {processed_files} files (valid data from {len(df['文件名称'].unique())} files) to {output_file}")
    else:
        print(f"No data extracted from {processed_files} files scanned.")

if __name__ == "__main__":
    main()
